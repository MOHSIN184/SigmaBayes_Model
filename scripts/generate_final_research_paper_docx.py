from pathlib import Path
from typing import Iterable

import pandas as pd

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install it with: pip install python-docx"
    ) from exc


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = PROJECT_ROOT / "results"
TABLE_DIR = RESULTS_DIR / "tables"
FIG_DIR = RESULTS_DIR / "figures"
OUTPUT_PATH = PROJECT_ROOT / "Final_Research_Paper.docx"


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def safe_float(value, default=float("nan")):
    try:
        return float(value)
    except Exception:
        return default


def fmt(value, ndigits=4):
    v = safe_float(value)
    if pd.isna(v):
        return "N/A"
    return f"{v:.{ndigits}f}"


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, lines: Iterable[str]):
    for line in lines:
        doc.add_paragraph(line)


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 20):
    doc.add_paragraph(title)
    if df.empty:
        doc.add_paragraph("Table unavailable in workspace outputs.")
        return

    shown = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.style = "Table Grid"
    for i, col in enumerate(shown.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(shown.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = fmt(val, ndigits=4)
            else:
                cells[i].text = str(val)


def add_figure(doc: Document, fig_counter: int, filename: str, caption: str, width=6.2):
    fig_path = FIG_DIR / filename
    if not fig_path.exists():
        doc.add_paragraph(f"Figure {fig_counter} unavailable: {filename}")
        return fig_counter + 1

    doc.add_picture(str(fig_path), width=Inches(width))
    cap = doc.add_paragraph(f"Figure {fig_counter}. {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return fig_counter + 1


def get_metric(df: pd.DataFrame, experiment: str, metric: str):
    if df.empty:
        return float("nan")
    rows = df[df["experiment"] == experiment]
    if rows.empty or metric not in rows.columns:
        return float("nan")
    return safe_float(rows.iloc[0][metric])


def main():
    dataset_summary = read_csv(TABLE_DIR / "dataset_summary.csv")
    class_dist = read_csv(TABLE_DIR / "dataset_class_distribution.csv")
    final_metrics = read_csv(TABLE_DIR / "final_metrics_summary.csv")
    model_cmp = read_csv(TABLE_DIR / "final_model_comparison_with_hybrid.csv")
    reliability_cmp = read_csv(TABLE_DIR / "final_reliability_comparison.csv")
    cv_summary = read_csv(TABLE_DIR / "cv_summary_wide.csv")
    sigma_diag = read_csv(TABLE_DIR / "analysis_sigma_class_diagnostic.csv")
    sigma_confusions = read_csv(TABLE_DIR / "analysis_sigma_confusion_pairs.csv")

    binary_acc = get_metric(final_metrics, "binary_cnn", "accuracy")
    binary_f1 = get_metric(final_metrics, "binary_cnn", "f1_macro")
    binary_auc = get_metric(final_metrics, "binary_cnn", "auroc")
    sigma_acc = get_metric(final_metrics, "safe_f1_selected_sigma_cnn", "accuracy")
    sigma_f1 = get_metric(final_metrics, "safe_f1_selected_sigma_cnn", "f1_macro")
    sigma_auc = get_metric(final_metrics, "safe_f1_selected_sigma_cnn", "auroc_macro_ovr")

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    title = doc.add_paragraph(
        "BayesSigma: Reliability-Aware Promoter and Sigma-Factor Prediction "
        "Using Calibrated Deep Learning, Classical Baselines, and Conformal Inference"
    )
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Generated from project evidence in the BayesSigma workspace")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Candidate Publication Titles", level=1)
    add_paragraphs(
        doc,
        [
            "1) BayesSigma: A Reliability-Aware Benchmark for Bacterial Promoter and Sigma-Factor Prediction",
            "2) Trustworthy Bacterial Promoter Modeling with Calibration, MC Dropout, and Conformal Prediction",
            "3) Beyond Accuracy in Promoter Classification: A Comparative Reliability Study of CNN and k-mer Models",
        ],
    )

    add_heading(doc, "Abstract", level=1)
    add_paragraphs(
        doc,
        [
            "Promoter recognition and sigma-factor assignment are central tasks in bacterial genomics, but classical accuracy-only reporting can mask model overconfidence and minority-class failure. "
            "This study presents BayesSigma, an evidence-based reliability framework implemented in Python/PyTorch and scikit-learn for (i) binary promoter classification and (ii) six-class sigma-factor classification (Sigma24/28/32/38/54/70). "
            "Using curated 81-bp sequences (binary: 5,258 train and 1,315 test; sigma: 2,315 train and 583 test), we evaluated one-hot CNN models, multiple sigma training variants, k-mer classical baselines, and hybrid CNN+k-mer models. "
            "Reliability was quantified with temperature scaling, Expected Calibration Error (ECE), Brier score, MC Dropout uncertainty, and split conformal prediction at 90% and 95% target confidence. "
            f"The final binary CNN reached accuracy={fmt(binary_acc)}, macro-F1={fmt(binary_f1)}, AUROC={fmt(binary_auc)}; however, calibrated 5-mer Random Forest provided stronger binary endpoint performance in comparison tables. "
            f"The selected sigma-safe CNN reached accuracy={fmt(sigma_acc)}, macro-F1={fmt(sigma_f1)}, macro-AUROC(OVR)={fmt(sigma_auc)}, while sigma macro-F1 remained modest across models and cross-validation. "
            "Conformal coverage targets were achieved but sigma prediction sets were large, indicating meaningful uncertainty under class overlap and imbalance. "
            "These findings support a publication framing as a reliability-aware benchmark rather than a state-of-the-art sigma classifier.",
        ],
    )

    add_heading(doc, "Keywords", level=1)
    doc.add_paragraph(
        "Promoter prediction; sigma factor classification; uncertainty quantification; calibration; conformal prediction; macro-F1; reliability benchmarking; bacterial genomics"
    )

    add_heading(doc, "1. Introduction", level=1)
    add_paragraphs(
        doc,
        [
            "Promoter and sigma-factor prediction often operates in high-risk settings where false certainty can be as damaging as raw error. BayesSigma was designed to answer not only which label is predicted, but also how trustworthy that prediction is.",
            "The repository reveals a deliberate shift from pure CNN performance optimization toward reliability-aware comparative modeling: neural models, classical k-mer baselines, cross-validation stability checks, calibration diagnostics, uncertainty estimation, and conformal prediction sets.",
            "The primary research gap addressed here is transparency of confidence in promoter classification workflows. The project contribution is practical: an end-to-end, reproducible reliability pipeline with explicit failure-mode reporting, especially for imbalanced sigma classes.",
        ],
    )

    add_heading(doc, "2. Literature Review Framework", level=1)
    add_paragraphs(
        doc,
        [
            "Based on project structure and analysis artifacts, BayesSigma aligns with four methodological streams: (i) sequence-based promoter classifiers (k-mer and CNN), (ii) imbalance-aware multiclass learning, (iii) probability calibration and reliability diagrams, and (iv) conformal uncertainty sets.",
            "The implemented comparison logic suggests an explicit critique of single-model claims: deep learning is tested, but not assumed superior. The repository therefore supports a benchmark-style narrative centered on reliability and robustness rather than headline accuracy.",
            "A full citation-level review is not present in the workspace outputs; thus, this section is framed as a methodological map to position the proposed approach in current research trends.",
        ],
    )

    add_heading(doc, "3. Materials and Methods", level=1)
    add_heading(doc, "3.1 Datasets", level=2)
    add_paragraphs(
        doc,
        [
            "All sequences are fixed length (81 bp), with no invalid DNA symbols detected in generated checks.",
            "Binary task: Non-Promoter vs Promoter. Sigma task: Sigma24, Sigma28, Sigma32, Sigma38, Sigma54, Sigma70.",
            "FASTA parsing and label mapping are implemented in src/data_loader.py with stratified splitting into train/validation/calibration subsets.",
        ],
    )
    add_table_from_df(doc, dataset_summary, "Table 1. Dataset summary from results/tables/dataset_summary.csv")
    add_table_from_df(doc, class_dist, "Table 2. Class distribution percentages")

    add_heading(doc, "3.2 Data Processing Pipeline", level=2)
    add_paragraphs(
        doc,
        [
            "Pipeline reconstruction from scripts and source code:",
            "1) FASTA ingestion and sequence normalization to uppercase.",
            "2) Quality checks for invalid symbols and sequence lengths.",
            "3) Task-specific labeling (binary or sigma-class mapping).",
            "4) Stratified split of original train into train/validation/calibration.",
            "5) Feature encoding: one-hot tensor (4 x 81) for CNN; normalized k-mer frequencies for classical models.",
            "6) Model training with class-weighted loss for imbalance-sensitive tasks.",
            "7) Untouched test evaluation.",
            "8) Post-hoc reliability analyses: MC Dropout, temperature scaling, conformal prediction.",
        ],
    )

    add_heading(doc, "3.3 Model Architectures and Learning", level=2)
    add_paragraphs(
        doc,
        [
            "CNN backbone (src/models.py): Conv1D(4->64, k=7), Conv1D(64->128, k=5), Conv1D(128->256, k=3), batch normalization, max pooling, adaptive pooling, MLP head with dropout.",
            "Sigma variants: original class-weighted CNN; improved WeightedRandomSampler+class-weighting configuration; safe model with lower learning rate and validation macro-F1 model selection.",
            "Classical baselines: Logistic Regression, Linear SVM, Random Forest on 3-mer/4-mer/5-mer frequency vectors.",
            "Hybrid model: fused CNN branch + k-mer MLP branch (src/hybrid_models.py).",
        ],
    )

    add_heading(doc, "3.4 Mathematical Formulation", level=2)
    add_paragraphs(
        doc,
        [
            "For logits z and temperature T, calibrated probabilities follow p_i = softmax(z_i / T).",
            "Class-weighted cross-entropy used in CNN training: L = - sum_c w_c y_c log(p_c).",
            "Expected calibration error over bins B_m: ECE = sum_m (|B_m|/n) * |acc(B_m)-conf(B_m)|.",
            "Conformal nonconformity score for true class y: s = 1 - p(y|x). Quantile q_hat is computed from calibration scores, and test prediction set is C(x) = {k : p(k|x) >= 1-q_hat}.",
        ],
    )

    add_heading(doc, "4. Experimental Setup", level=1)
    add_paragraphs(
        doc,
        [
            "Software stack (requirements.txt): numpy, pandas, scikit-learn, matplotlib, torch, tqdm, jupyter.",
            "Reproducibility controls: fixed seeds in training utilities and deterministic cuDNN mode when available.",
            "Representative settings from scripts: binary CNN (30 epochs, lr=1e-3, patience=8), original sigma CNN (40 epochs, lr=1e-3), improved sigma sampler variant (80 epochs, lr=5e-4), safe sigma CNN (100 epochs, lr=3e-4, selection by validation macro-F1).",
            "Hardware details are not explicitly logged in result tables, representing a reproducibility metadata gap.",
        ],
    )

    add_heading(doc, "5. Results and Discussion", level=1)
    add_table_from_df(doc, final_metrics, "Table 3. Final CNN-based metrics summary")
    add_table_from_df(doc, model_cmp, "Table 4. CNN vs k-mer vs hybrid comparison")
    add_table_from_df(doc, reliability_cmp, "Table 5. Reliability comparison (calibration + conformal)")

    fig = 1
    fig = add_figure(
        doc,
        fig,
        "binary_roc_curve.png",
        "Binary ROC curve for the CNN model, showing discrimination strength on untouched test data.",
    )
    fig = add_figure(
        doc,
        fig,
        "binary_reliability_after.png",
        "Binary reliability diagram after temperature scaling.",
    )
    fig = add_figure(
        doc,
        fig,
        "sigma_safe_confusion_matrix.png",
        "Confusion matrix for the selected safe sigma CNN model.",
    )
    fig = add_figure(
        doc,
        fig,
        "sigma_safe_uncertainty_by_class.png",
        "Class-wise uncertainty statistics for sigma prediction.",
    )
    fig = add_figure(
        doc,
        fig,
        "analysis_sigma_class_imbalance.png",
        "Sigma class imbalance profile, with minority classes highlighted.",
    )
    fig = add_figure(
        doc,
        fig,
        "analysis_sigma_confusion_pair_heatmap.png",
        "Dominant sigma confusion patterns across true-predicted label pairs.",
    )
    fig = add_figure(
        doc,
        fig,
        "analysis_sigma_f1_vs_uncertainty.png",
        "Relationship between per-class F1 and entropy-based uncertainty in sigma prediction.",
    )
    fig = add_figure(
        doc,
        fig,
        "cv_sigma_model_comparison.png",
        "5-fold cross-validation macro-F1 comparison for sigma models.",
    )

    add_paragraphs(
        doc,
        [
            f"Binary endpoint: final CNN yielded accuracy={fmt(binary_acc)}, macro-F1={fmt(binary_f1)}, AUROC={fmt(binary_auc)}; however, comparison tables show that calibrated 5-mer Random Forest is stronger for binary deployment-oriented performance and calibration.",
            f"Sigma endpoint: selected safe CNN achieved accuracy={fmt(sigma_acc)} and macro-F1={fmt(sigma_f1)} with macro-AUROC(OVR)={fmt(sigma_auc)}, but macro-F1 remained below 0.50 across methods and cross-validation summaries.",
            "The improved sampler sigma variant underperformed, indicating over-correction when combining weighted sampling with class-weighted loss.",
            "Conformal results reached near-target coverage but produced large sigma prediction sets, which is consistent with intrinsic ambiguity and overlap among sigma-promoter patterns.",
        ],
    )

    add_heading(doc, "6. Comparative Analysis", level=1)
    add_paragraphs(
        doc,
        [
            "Binary: k-mer Random Forest baseline outperformed CNN and hybrid variants on key summary endpoints in final comparison tables.",
            "Sigma: performance ordering is split-sensitive, but all configurations exhibit modest macro-F1; neither hybrid fusion nor aggressive resampling solved minority-class fragility.",
            "Cross-validation confirms the same qualitative conclusion: binary is stable and usable; sigma remains difficult and should be interpreted cautiously.",
        ],
    )
    add_table_from_df(doc, cv_summary, "Table 6. Cross-validation summary (mean and standard deviation)")

    add_heading(doc, "7. Research Contributions", level=1)
    add_paragraphs(
        doc,
        [
            "1) End-to-end reliability-aware promoter/sigma pipeline combining prediction, calibration, uncertainty, and conformal sets.",
            "2) Transparent multi-family comparison across CNN, classical k-mer baselines, and hybrid fusion.",
            "3) Explicit failure-mode documentation for sigma classes via confusion, uncertainty, and class-diagnostic analyses.",
            "4) Reproducible scripts generating publication-oriented tables and figures.",
        ],
    )

    add_heading(doc, "8. Limitations", level=1)
    add_paragraphs(
        doc,
        [
            "Single dataset family and no cross-species external validation in current artifacts.",
            "No wet-lab validation evidence in the workspace.",
            "Severe sigma imbalance (especially Sigma54, Sigma28, Sigma38) constrains balanced multiclass learning.",
            "Hardware/runtime metadata are not fully captured in result tables.",
        ],
    )

    add_heading(doc, "9. Future Work", level=1)
    add_paragraphs(
        doc,
        [
            "Add external and cross-species promoter datasets to quantify generalization.",
            "Expand minority sigma classes and include motif-level biological interpretation workflows.",
            "Report confidence-conditioned operating points for deployment thresholds.",
            "Integrate standardized experiment tracking (hardware, runtime, package versions, seed registry).",
        ],
    )

    add_heading(doc, "10. Conclusion", level=1)
    add_paragraphs(
        doc,
        [
            "BayesSigma demonstrates that reliability-centric evaluation materially changes conclusions in promoter modeling. The binary task is robust and operationally promising under calibrated/classical modeling, while sigma-factor assignment remains challenging despite multiple neural and hybrid strategies. "
            "The strongest evidence in this repository supports publication as a reliability-aware benchmark and methodological framework, not as a high-performance sigma-factor classifier claim.",
        ],
    )

    add_heading(doc, "Appendix: Additional Evidence Tables", level=1)
    add_table_from_df(doc, sigma_diag, "Table 7. Sigma class diagnostic summary", max_rows=10)
    add_table_from_df(doc, sigma_confusions, "Table 8. Top sigma confusion pairs", max_rows=20)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
